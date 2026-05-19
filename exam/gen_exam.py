"""
Génère examen_<slug>.docx et dossier_<slug>.docx à partir d'un fichier de config JSON.

Usage :
    python3 exam/gen_exam.py              # utilise exam_p1.json par défaut
    python3 exam/gen_exam.py exam_p2.json # autre config

Config JSON :
    {
      "titre_examen":  "Examen — Période 1 : ...",
      "titre_dossier": "Dossier documentaire — ...",
      "total_points":  20,
      "questions":     ["Q2", "Q33", ...]   // IDs dans l'ordre voulu
    }

Les DOCX générés sont exclus du dépôt (.gitignore).
"""
import json, re, subprocess, sys
from pathlib import Path
from lxml import etree
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

EXAM_DIR = Path(__file__).parent
REPO_DIR = EXAM_DIR.parent
IMG_DIR  = REPO_DIR / 'images'

# ── CHARGEMENT CONFIG + DONNÉES ──────────────────────────────────────────────
config_file = EXAM_DIR / (sys.argv[1] if len(sys.argv) > 1 else 'exam_p1.json')
with open(config_file, encoding='utf-8') as f:
    config = json.load(f)

slug = config_file.stem  # ex. "exam_p1"

# Extraction Node.js : lit questions.js et retourne les questions demandées
_node_script = f"""
const fs = require('fs');
const src = fs.readFileSync({json.dumps(str(REPO_DIR / 'questions.js'))}, 'utf8');
const {{ REGLETTES, IMAGE_DB, QUESTIONS }} = new Function(src + '; return {{ REGLETTES, IMAGE_DB, QUESTIONS }};')();
const ids = {json.dumps(config['questions'])};
const result = {{ questions: [], IMAGE_DB }};
ids.forEach(id => {{
  const q = QUESTIONS.find(q => q.id === id);
  if (!q) {{ process.stderr.write('Question introuvable : ' + id + '\\n'); process.exit(1); }}
  const r = (REGLETTES && REGLETTES[id]) ? REGLETTES[id] : null;
  result.questions.push({{ q, r }});
}});
process.stdout.write(JSON.stringify(result));
"""
_res = subprocess.run(['node', '-e', _node_script], capture_output=True, text=True)
if _res.returncode != 0:
    print('Erreur Node.js :', _res.stderr); sys.exit(1)

data      = json.loads(_res.stdout)
questions = data['questions']
IMAGE_DB  = data['IMAGE_DB']

# ── DOCUMENT LETTER MAP ──────────────────────────────────────────────────────
_global_letter = 0
doc_map      = []
all_docs_flat = []

for qidx, item in enumerate(questions):
    q        = item['q']
    q_letters = []
    for d in q.get('documents', []):
        for col in d.get('cols', [d]):
            letter = chr(65 + _global_letter)
            q_letters.append({'d': d, 'col': col, 'letter': letter})
            all_docs_flat.append({'qnum': qidx+1, 'q': q, 'col': col, 'letter': letter})
            _global_letter += 1
    doc_map.append(q_letters)

# ── REMAP ÉNONCÉ ─────────────────────────────────────────────────────────────
def remap_enonce(text, q_letters):
    lmap = {chr(65+i): e['letter'] for i, e in enumerate(q_letters)}
    def replace_doc(m):
        return m.group(1) + ''.join(lmap.get(c, c) if c.isupper() else c for c in m.group(2))
    return re.sub(r'(documents?\s+)([A-Z](?:\s*(?:et|à|,)\s*[A-Z])*)', replace_doc, text, flags=re.IGNORECASE)

# ── FONT / PARAGRAPH HELPERS ─────────────────────────────────────────────────
def set_font(run, size=11, bold=False, italic=False, color=None):
    run.font.name = 'Aptos'; run.font.size = Pt(size)
    run.bold = bold; run.italic = italic
    if color: run.font.color.rgb = RGBColor(*color)

def add_para(doc_or_cell, text='', size=11, bold=False, italic=False,
             color=None, align=None, sb=0, sa=6):
    p = doc_or_cell.add_paragraph()
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    if align: p.alignment = align
    if text:
        run = p.add_run(text); set_font(run, size=size, bold=bold, italic=italic, color=color)
    return p

def add_formatted(doc_or_cell, text, size=10.5, sb=2, sa=4):
    p = doc_or_cell.add_paragraph()
    p.paragraph_format.space_before = Pt(sb); p.paragraph_format.space_after = Pt(sa)
    for part in re.split(r'(\*\*.*?\*\*)', text):
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2]); set_font(run, size=size, bold=True)
        else:
            run = p.add_run(part); set_font(run, size=size)
    return p

def add_image_para(container, img_path, width_cm=14.5):
    p = container.add_paragraph()
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
    try:    p.add_run().add_picture(str(img_path), width=Cm(width_cm))
    except: p.add_run(f'[Image: {img_path.name}]').italic = True
    return p

# ── BORDER / TABLE HELPERS ────────────────────────────────────────────────────
def set_cell_borders(cell, top=None, bottom=None, left=None, right=None, sz=4, color='000000'):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    for ex in tcPr.findall(qn('w:tcBorders')): tcPr.remove(ex)
    tcB = OxmlElement('w:tcBorders')
    for side, show in [('top',top),('bottom',bottom),('left',left),('right',right),
                       ('insideH',None),('insideV',None)]:
        el = OxmlElement(f'w:{side}')
        if show:
            el.set(qn('w:val'),'single'); el.set(qn('w:sz'),str(sz)); el.set(qn('w:color'),color)
        else:
            el.set(qn('w:val'),'nil')
        tcB.append(el)
    tcPr.append(tcB)

def set_cell_no_border(cell):   set_cell_borders(cell)
def set_cell_all_borders(cell, sz=4, color='000000'):
    set_cell_borders(cell, top=True, bottom=True, left=True, right=True, sz=sz, color=color)

def set_table_full_width(tbl):
    tbl_pr = tbl._tblPr
    for el in list(tbl_pr):
        if el.tag == qn('w:tblW'): tbl_pr.remove(el)
    w = OxmlElement('w:tblW'); w.set(qn('w:w'),'5000'); w.set(qn('w:type'),'pct')
    tbl_pr.append(w)

def set_table_no_border(tbl):
    tbl_pr = tbl._tblPr
    for el in list(tbl_pr):
        if el.tag == qn('w:tblBorders'): tbl_pr.remove(el)
    tb = OxmlElement('w:tblBorders')
    for side in ('top','left','bottom','right','insideH','insideV'):
        el = OxmlElement(f'w:{side}'); el.set(qn('w:val'),'nil'); tb.append(el)
    tbl_pr.append(tb)

# ── ELLIPSE DRAWINGML ─────────────────────────────────────────────────────────
_shape_id = 2000
def add_ellipse_to_para(para, cx_emu=450000, cy_emu=450000):
    global _shape_id; _shape_id += 1; sid = str(_shape_id)
    W='http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    WP='http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
    A='http://schemas.openxmlformats.org/drawingml/2006/main'
    WPS='http://schemas.microsoft.com/office/word/2010/wordprocessingShape'
    xml = f'<w:r xmlns:w="{W}" xmlns:wp="{WP}" xmlns:a="{A}" xmlns:wps="{WPS}"><w:drawing><wp:inline distT="0" distB="0" distL="0" distR="0"><wp:extent cx="{cx_emu}" cy="{cy_emu}"/><wp:effectExtent l="0" t="0" r="0" b="0"/><wp:docPr id="{sid}" name="Ellipse {sid}"/><wp:cNvGraphicFramePr/><a:graphic><a:graphicData uri="{WPS.replace("Shape","Shape")}"><wps:wsp><wps:cNvSpPr/><wps:spPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="{cx_emu}" cy="{cy_emu}"/></a:xfrm><a:prstGeom prst="ellipse"><a:avLst/></a:prstGeom><a:noFill/><a:ln><a:solidFill><a:srgbClr val="000000"/></a:solidFill></a:ln></wps:spPr><wps:bodyPr anchor="ctr"/></wps:wsp></a:graphicData></a:graphic></wp:inline></w:drawing></w:r>'
    para._p.append(etree.fromstring(xml))

# ── ESPACES RÉPONSE ───────────────────────────────────────────────────────────
def add_answer_lines(doc, n, height_pt=28):
    tbl = doc.add_table(rows=n, cols=1)
    tbl.style = 'Normal Table'; set_table_full_width(tbl); set_table_no_border(tbl)
    for row in tbl.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY; row.height = Pt(height_pt)
        cell = row.cells[0]; set_cell_borders(cell, bottom=True)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)

def add_response_image(doc, ref):
    add_image_para(doc, IMG_DIR / ref, width_cm=14.5)

def add_response_tableau_2col(doc, n_rows=4):
    tbl = doc.add_table(rows=n_rows+1, cols=2)
    tbl.style = 'Normal Table'; set_table_full_width(tbl)
    for row in tbl.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY; row.height = Pt(22)
        for cell in row.cells:
            set_cell_all_borders(cell, sz=4, color='999999')
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)

def add_response_mettre_en_relation(doc, elements, n_rows=5):
    cols = len(elements)
    tbl = doc.add_table(rows=n_rows+1, cols=cols)
    tbl.style = 'Normal Table'; set_table_full_width(tbl)
    for j, elem in enumerate(elements):
        cell = tbl.rows[0].cells[j]; set_cell_all_borders(cell, sz=4, color='999999')
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
        run = p.add_run(elem); set_font(run, size=9, bold=True)
    for i in range(1, n_rows+1):
        tbl.rows[i].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY; tbl.rows[i].height = Pt(22)
        for j in range(cols):
            cell = tbl.rows[i].cells[j]; set_cell_all_borders(cell, sz=4, color='999999')
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)

def add_response_situer_dans_lespace(doc, elements):
    els = elements if len(elements) == 2 else ['Élément 1','Élément 2']
    tbl = doc.add_table(rows=2, cols=2); tbl.style = 'Normal Table'; tbl.autofit = True
    for j, label in enumerate(els):
        cell = tbl.rows[0].cells[j]; set_cell_all_borders(cell)
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
        run = p.add_run(label); set_font(run, size=10, bold=True)
    tbl.rows[1].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST; tbl.rows[1].height = Pt(50)
    for j in range(2):
        cell = tbl.rows[1].cells[j]; set_cell_all_borders(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)
        add_ellipse_to_para(p)

# ── RÉGLETTES ────────────────────────────────────────────────────────────────
def rcell(cell, text='', size=8, bold=False, halign=WD_ALIGN_PARAGRAPH.CENTER):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]; p.alignment = halign
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
    if text: run = p.add_run(text); set_font(run, size=size, bold=bold)
    return p

def rcell_add(cell, text, size=8, halign=WD_ALIGN_PARAGRAPH.CENTER):
    p = cell.add_paragraph(); p.alignment = halign
    p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text); set_font(run, size=size)

def add_reglette_standard(doc, r):
    niveaux = r.get('niveaux', [])
    if not niveaux: return
    n = len(niveaux)
    tbl = doc.add_table(rows=1, cols=n+1)
    tbl.style = 'Table Grid'; tbl.alignment = WD_TABLE_ALIGNMENT.CENTER; set_table_full_width(tbl)
    rcell(tbl.rows[0].cells[0], r.get('oi',''), size=8, bold=True)
    set_cell_all_borders(tbl.rows[0].cells[0])
    for i, niv in enumerate(niveaux):
        cell = tbl.rows[0].cells[i+1]
        rcell(cell, f"{niv['pts']} pt{'s' if niv['pts']>1 else ''}", size=8, bold=True)
        rcell_add(cell, niv.get('desc',''), size=7.5)
        set_cell_all_borders(cell)

def add_reglette_acteur_positions(doc, r):
    oi = r.get('oi','')
    tbl = doc.add_table(rows=5, cols=4)
    tbl.style = 'Table Grid'; tbl.alignment = WD_TABLE_ALIGNMENT.CENTER; set_table_full_width(tbl)
    tbl.cell(0,0).merge(tbl.cell(4,0)); rcell(tbl.cell(0,0), oi, size=8, bold=True)
    tbl.cell(0,1).merge(tbl.cell(3,1))
    rcell(tbl.cell(0,1), "L'élève nomme correctement l'acteur qui présente une position différente", size=8, halign=WD_ALIGN_PARAGRAPH.LEFT)
    descs = ["et présente correctement les deux positions.",
             "et présente correctement une position et plus ou moins correctement l'autre position.",
             "et présente plus ou moins correctement les deux positions, ou présente correctement une position et incorrectement l'autre ou ne la présente pas.",
             "et présente tout au plus une seule position plus ou moins correctement."]
    for i, (desc, pts) in enumerate(zip(descs, ["3 points","2 points","1 point","0 point"])):
        rcell(tbl.cell(i,2), desc, size=7.5, halign=WD_ALIGN_PARAGRAPH.LEFT)
        rcell(tbl.cell(i,3), pts, size=8, bold=True)
    tbl.cell(4,1).merge(tbl.cell(4,2))
    rcell(tbl.cell(4,1), "L'élève nomme incorrectement l'acteur qui présente une position différente ou ne le nomme pas.", size=7.5, halign=WD_ALIGN_PARAGRAPH.LEFT)
    rcell(tbl.cell(4,3), "0 point", size=8, bold=True)

def add_reglette_3el_2liens(doc, r):
    oi = r.get('oi','')
    tbl = doc.add_table(rows=6, cols=4)
    tbl.style = 'Table Grid'; tbl.alignment = WD_TABLE_ALIGNMENT.CENTER; set_table_full_width(tbl)
    tbl.cell(0,0).merge(tbl.cell(5,0)); rcell(tbl.cell(0,0), oi, size=8, bold=True)
    tbl.cell(0,1).merge(tbl.cell(2,1))
    rcell(tbl.cell(0,1), "L'élève précise les trois éléments", size=8, halign=WD_ALIGN_PARAGRAPH.LEFT)
    for i, (desc, pts) in enumerate([("et établit correctement deux liens de causalité.","3 points"),
                                      ("et établit correctement un lien de causalité.","2 points"),
                                      ("mais n'établit correctement aucun lien de causalité.","1 point")]):
        rcell(tbl.cell(i,2), desc, size=7.5, halign=WD_ALIGN_PARAGRAPH.LEFT)
        rcell(tbl.cell(i,3), pts, size=8, bold=True)
    tbl.cell(3,1).merge(tbl.cell(4,1))
    rcell(tbl.cell(3,1), "L'élève précise deux éléments", size=8, halign=WD_ALIGN_PARAGRAPH.LEFT)
    for i, (desc, pts) in enumerate([("et établit correctement un lien de causalité.","2 points"),
                                      ("mais n'établit correctement aucun lien de causalité.","1 point")]):
        rcell(tbl.cell(3+i,2), desc, size=7.5, halign=WD_ALIGN_PARAGRAPH.LEFT)
        rcell(tbl.cell(3+i,3), pts, size=8, bold=True)
    tbl.cell(5,1).merge(tbl.cell(5,2))
    rcell(tbl.cell(5,1), "L'élève précise un seul élément ou n'en précise pas.", size=7.5, halign=WD_ALIGN_PARAGRAPH.LEFT)
    rcell(tbl.cell(5,3), "0 point", size=8, bold=True)

def add_reglette(doc, r):
    if not r: return
    v = r.get('variante','')
    add_para(doc, '', sb=6, sa=2)
    if   v == 'acteur-positions':    add_reglette_acteur_positions(doc, r)
    elif v == '3 éléments — 2 liens': add_reglette_3el_2liens(doc, r)
    else:                              add_reglette_standard(doc, r)

# ── GÉNÉRATION EXAMEN ─────────────────────────────────────────────────────────
def gen_examen():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.59); sec.page_height = Cm(27.94)
    sec.left_margin = sec.right_margin = Cm(2.54)
    sec.top_margin  = sec.bottom_margin = Cm(2.54)

    add_para(doc, config['titre_examen'], size=14, bold=True, sb=0, sa=2, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, f"Total : {config['total_points']} points", size=11, italic=True, sb=0, sa=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    for qnum, item in enumerate(questions, 1):
        q = item['q']; r = item['r']
        q_letters = doc_map[qnum-1]
        enonce = remap_enonce(q.get('enonce',''), q_letters)

        add_para(doc, f'Question {qnum}', size=11, bold=True, color=(30,30,30), sb=12, sa=3)
        add_formatted(doc, enonce, size=10.5, sb=0, sa=4)

        rep   = q.get('reponse') or {}
        rtype = rep.get('type','')
        if   rtype == 'lignes':               add_answer_lines(doc, rep.get('nombre', 2))
        elif rtype == 'image':                add_response_image(doc, rep.get('ref',''))
        elif rtype == 'tableau_2col':         add_response_tableau_2col(doc)
        elif rtype == 'mettre-en-relation':   add_response_mettre_en_relation(doc, rep.get('elements',[]))
        elif rtype == 'situer-dans-lespace':  add_response_situer_dans_lespace(doc, rep.get('elements',[]))

        add_reglette(doc, r)
        add_para(doc, '', sb=4, sa=0)

    out = EXAM_DIR / f'examen_{slug.replace("exam_","")}.docx'
    doc.save(str(out)); print(f'Examen → {out.name}')

# ── GÉNÉRATION DOSSIER ────────────────────────────────────────────────────────
def fill_doc_cell(cell, col, letter, width_cm=7.0):
    titre = re.sub(r'Document [A-Z]', f'Document {letter}', col.get('titre','Document '+letter), count=1)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(1)
    run = p.add_run(titre); set_font(run, size=10, bold=True, color=(30,30,100))
    if col.get('soustitre'):
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_before = Pt(0); p2.paragraph_format.space_after = Pt(2)
        run2 = p2.add_run(col['soustitre']); set_font(run2, size=9, italic=True, color=(80,80,80))
    if col.get('ref'):
        p_img = cell.add_paragraph()
        p_img.paragraph_format.space_before = Pt(3); p_img.paragraph_format.space_after = Pt(3)
        try:    p_img.add_run().add_picture(str(IMG_DIR / col['ref']), width=Cm(width_cm))
        except: p_img.add_run(f'[Image: {col["ref"]}]').italic = True
    elif col.get('texte'):
        texte = col['texte']
        if not texte.startswith('«'): texte = f'« {texte} »'
        add_formatted(cell, texte, size=9.5, sb=2, sa=3)
    meta = ' — '.join(filter(None, [col.get('auteur',''), col.get('source','')]))
    if meta:
        pm = cell.add_paragraph()
        pm.paragraph_format.space_before = Pt(1); pm.paragraph_format.space_after = Pt(3)
        run_m = pm.add_run(meta); set_font(run_m, size=8, italic=True, color=(100,100,100))

def gen_dossier():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.59); sec.page_height = Cm(27.94)
    sec.left_margin = sec.right_margin = Cm(2.54)
    sec.top_margin  = sec.bottom_margin = Cm(2.54)

    add_para(doc, config['titre_dossier'], size=14, bold=True, sb=0, sa=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    i = 0
    while i < len(all_docs_flat):
        left  = all_docs_flat[i]
        right = all_docs_flat[i+1] if i+1 < len(all_docs_flat) else None
        tbl   = doc.add_table(rows=1, cols=2)
        tbl.style = 'Table Grid'; tbl.alignment = WD_TABLE_ALIGNMENT.LEFT; set_table_full_width(tbl)
        cells = tbl.rows[0].cells
        fill_doc_cell(cells[0], left['col'], left['letter'])
        set_cell_all_borders(cells[0], color='CCCCCC')
        if right:
            fill_doc_cell(cells[1], right['col'], right['letter'])
            set_cell_all_borders(cells[1], color='CCCCCC')
        else:
            set_cell_no_border(cells[1])
        add_para(doc, '', sb=4, sa=4)
        i += 2

    out = EXAM_DIR / f'dossier_{slug.replace("exam_","")}.docx'
    doc.save(str(out)); print(f'Dossier → {out.name}')

# ── MAIN ──────────────────────────────────────────────────────────────────────
if __name__ == '__main__':
    gen_examen()
    gen_dossier()
    print('Terminé.')
